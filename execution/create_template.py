from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contract_template():
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONTRATO SOCIAL DE CONSTITUIÇÃO DE\nSOCIEDADE EMPRESÁRIA LIMITADA")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph() # Spacer

    # Preamble (Multi-partner loop)
    p = doc.add_paragraph()
    p.add_run("Pelo presente instrumento particular de Contrato Social:\n")
    
    # Partner loop using simple jinja (no 'tr' tag)
    p.add_run("{% for p in partners %}")
    p.add_run("{{ p.name }}").bold = True
    p.add_run(", {{ p.nationality }}, {{ p.civil_state }}, {{ p.regime }}, {{ p.profession }}, nascido(a) em {{ p.birth_date }}, nº do CPF {{ p.cpf }}, residente e domiciliada na {{ p.address }}")
    p.add_run("{% if not loop.last %};\n{% else %};{% endif %}")
    p.add_run("{% endfor %}")

    p = doc.add_paragraph()
    p.add_run("Resolvem, em comum acordo, constituir uma sociedade empresária limitada, nos termos da Lei n° 10.406/2002, mediante as condições e cláusulas seguintes:")

    # Clause 1 - Name
    doc.add_heading('CLÁUSULA I - DO NOME EMPRESARIAL', level=2)
    p = doc.add_paragraph()
    p.add_run("A sociedade adotará como nome empresarial: ")
    p.add_run("{{ company_name }}").bold = True
    p.add_run(".")

    # Clause 2 - Head office
    doc.add_heading('CLÁUSULA II - DA SEDE', level=2)
    p = doc.add_paragraph("A sociedade terá sua sede no seguinte endereço: ")
    p.add_run("{{ company_address }}")
    p.add_run(".")

    # Clause 3 - Object
    doc.add_heading('CLÁUSULA III - DO OBJETO SOCIAL', level=2)
    p = doc.add_paragraph("A sociedade terá por objeto o exercício das seguintes atividades econômica: ")
    p.add_run("{{ company_object }}")
    p.add_run(".")
    
    # Activities List
    p = doc.add_paragraph("E exercerá as seguintes atividades:\n")
    p.add_run("{{ company_cnae_list }}")

    # Clause 4 - Duration
    doc.add_heading('CLÁUSULA IV - DO INÍCIO DAS ATIVIDADES E PRAZO DE DURAÇÃO', level=2)
    p = doc.add_paragraph("A sociedade iniciará suas atividades em {{ start_date }} e seu prazo de duração será por tempo indeterminado.")

    # Clause 5 - Capital
    doc.add_heading('CLÁUSULA V - DO CAPITAL', level=2)
    p = doc.add_paragraph("O capital será de {{ capital_currency }} ({{ capital_amount_text }}), dividido em {{ total_quotas }} quotas, no valor nominal de {{ quota_value }} cada uma, formado por {{ capital_currency }} em moeda corrente no Pais.")
    
    p = doc.add_paragraph("Parágrafo único. O capital encontra-se subscrito e integralizado pelos sócios da seguinte forma:")
    
    # Capital distribution as text list (docxtpl table row loops don't work when created via python-docx)
    p = doc.add_paragraph()
    p.add_run("{% for p in partners %}")
    p.add_run("• {{ p.name }}: {{ p.quotas }} quotas, {{ p.amount }}, {{ p.percent }}%")
    p.add_run("{% if not loop.last %}\n{% endif %}")
    p.add_run("{% endfor %}")
    
    p = doc.add_paragraph()
    p.add_run("TOTAL: {{ total_quotas }} quotas, {{ capital_currency }}, 100%").bold = True

    # Clause 6 - Administration
    doc.add_heading('CLÁUSULA VI - DA ADMINISTRAÇÃO', level=2)
    p = doc.add_paragraph("A administração da sociedade será exercida pelo(s) sócio(s) ")
    p.add_run("{{ administrator_names }}").bold = True
    p.add_run(" que representará(ão) legalmente a sociedade e poderá(ão) praticar todo e qualquer ato de gestão pertinente ao objeto social.")

    # Clause 15 - Forum
    doc.add_heading('CLÁUSULA XV - DO FORO', level=2)
    p = doc.add_paragraph("Fica eleito o Foro da Comarca de {{ forum_city }}, para qualquer ação fundada neste contrato.")

    # Signatures
    doc.add_paragraph()
    p = doc.add_paragraph("{{ forum_city }}, {{ signature_date }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Signature Loop
    p = doc.add_paragraph()
    p.add_run("{% for p in partners %}\n\n_______________________________________\n")
    p.add_run("{{ p.name }}\n")
    p.add_run("Sócio\n")
    p.add_run("{% endfor %}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('contract_template.docx')
    print("Template created: contract_template.docx")

if __name__ == "__main__":
    create_contract_template()
